import os, warnings
warnings.filterwarnings('ignore')

# Read DOCX
from docx import Document
folder = r'c:\Users\X1 Yoga\OneDrive\Desktop\New folder'
docx_files = [f for f in os.listdir(folder) if f.endswith('.docx')]
docx_path = os.path.join(folder, docx_files[0])
print(f"Reading: {docx_files[0]}")
doc = Document(docx_path)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f'P{i}: [{p.style.name}] {p.text[:300]}')
    if i > 150: break
print(f'\n=== Tables: {len(doc.tables)} ===')
for ti, table in enumerate(doc.tables):
    print(f'\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols')
    for ri, row in enumerate(table.rows[:8]):
        vals = [c.text[:50] for c in row.cells]
        print(f'  Row{ri}: {vals}')

# Read Excel deeper
import openpyxl
wb = openpyxl.load_workbook(os.path.join(folder, 'Data Vận hành.xlsx'), data_only=True)
ws = wb['Tất cả']
print('\n\n=== Excel: Tat ca sheet - Headers ===')
for i in range(1,5):
    for j in range(1, ws.max_column+1):
        v = ws.cell(i,j).value
        if v: print(f'  R{i}C{j}: {v}')

print('\n=== Sample rows ===')
for i in range(5, min(16, ws.max_row+1)):
    vals = []
    for j in range(1, ws.max_column+1):
        v = ws.cell(i,j).value
        if v: vals.append(f'C{j}={v}')
    if vals: print(f'  R{i}: {", ".join(vals)}')

# Stats
ams = set()
districts = set()
provinces = set()
buu_cucs = set()
for i in range(5, ws.max_row+1):
    am = ws.cell(i,20).value
    dist = ws.cell(i,19).value
    prov = ws.cell(i,18).value
    bc = ws.cell(i,3).value
    if am: ams.add(am)
    if dist: districts.add(dist)
    if prov: provinces.add(prov)
    if bc: buu_cucs.add(bc)

print(f'\nTotal rows: {ws.max_row-4}')
print(f'Unique AMs: {len(ams)} -> {sorted(ams)[:15]}')
print(f'Unique provinces: {len(provinces)} -> {sorted(provinces)}')
print(f'Unique districts: {len(districts)}')
print(f'Unique Buu cucs: {len(buu_cucs)}')
